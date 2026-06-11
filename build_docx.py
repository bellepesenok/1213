#!/usr/bin/env python3
"""Convert the bilingual markdown explanation into a formatted .docx file."""
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "国际贸易政策_解释_Объяснение.md"
OUT = "国际贸易政策_解释_Объяснение.docx"

doc = Document()

# Base style: a font with good CJK + Cyrillic coverage.
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(11)


def add_runs_with_bold(paragraph, text):
    """Split text on **bold** markers and add runs accordingly."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


with open(SRC, encoding="utf-8") as f:
    lines = f.readlines()

for raw in lines:
    line = raw.rstrip("\n")
    stripped = line.strip()

    if stripped == "" :
        continue
    if stripped == "---":
        continue

    # Title
    if line.startswith("# "):
        p = doc.add_heading(level=0)
        run = p.add_run(line[2:].strip())
        continue

    # Section heading per slide
    if line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=1)
        continue

    # Blockquote -> intro paragraph (italic, indented)
    if line.startswith(">"):
        text = line.lstrip(">").strip()
        if not text:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        r = p.add_run(text)
        r.italic = True
        continue

    # Bullet list (top level "- " or nested "  - ")
    m = re.match(r"^(\s*)-\s+(.*)$", line)
    if m:
        indent = len(m.group(1))
        content = m.group(2)
        style = "List Bullet 2" if indent >= 2 else "List Bullet"
        try:
            p = doc.add_paragraph(style=style)
        except KeyError:
            p = doc.add_paragraph(style="List Bullet")
        add_runs_with_bold(p, content)
        continue

    # Numbered list
    m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
    if m:
        content = m.group(3)
        p = doc.add_paragraph(style="List Number")
        add_runs_with_bold(p, content)
        continue

    # Normal paragraph
    p = doc.add_paragraph()
    add_runs_with_bold(p, stripped)

doc.save(OUT)
print("Saved", OUT)

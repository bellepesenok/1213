# -*- coding: utf-8 -*-
"""
生成《世界贸易组织概论》期末论文 Word 文档（含目录、脚注、图、表）。

题目：入世是"污染天堂"还是"技术外溢"？
      ——基于跨国面板数据对 WTO 成员身份与出口 CO2 强度关系的实证检验

输出：WTO期末论文_入世与出口碳强度.docx
依赖：python-docx；需先运行 generate_figures.py 生成 fig1-fig5.png。
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.packuri import PackURI
from docx.opc.part import Part

CN_FONT = "宋体"
CN_HEI = "黑体"
EN_FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)

FOOTNOTES_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
FOOTNOTES_RT = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"

_FOOTNOTES = []  # 收集脚注文本，最后写入 footnotes 部件


def set_run_font(run, cn=CN_FONT, en=EN_FONT, size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), en)
    rfonts.set(qn("w:hAnsi"), en)
    rfonts.set(qn("w:eastAsia"), cn)
    if color is not None:
        run.font.color.rgb = color


def add_footnote(paragraph, text):
    """在段落末尾追加一个脚注引用，并登记脚注文本。"""
    fid = len(_FOOTNOTES) + 1
    _FOOTNOTES.append(text)
    run = paragraph.add_run()
    rpr = run._element.get_or_add_rPr()
    rstyle = OxmlElement("w:rStyle"); rstyle.set(qn("w:val"), "FootnoteReference"); rpr.append(rstyle)
    valign = OxmlElement("w:vertAlign"); valign.set(qn("w:val"), "superscript"); rpr.append(valign)
    ref = OxmlElement("w:footnoteReference"); ref.set(qn("w:id"), str(fid))
    run._element.append(ref)
    return paragraph


def add_body(doc, text, size=12, first_indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_after=0, line=1.5, bold=False, cn=CN_FONT, footnote=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    if first_indent:
        pf.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    set_run_font(run, cn=cn, size=size, bold=bold)
    if footnote:
        add_footnote(p, footnote)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    pf = p.paragraph_format
    pf.space_before = Pt(12); pf.space_after = Pt(6)
    pf.line_spacing = 1.5; pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    run = p.add_run(text)
    set_run_font(run, cn=CN_HEI, size=14, bold=True, color=BLACK)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    pf = p.paragraph_format
    pf.space_before = Pt(6); pf.space_after = Pt(3)
    pf.line_spacing = 1.5; pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    run = p.add_run(text)
    set_run_font(run, cn=CN_HEI, size=12.5, bold=True, color=BLACK)
    return p


def add_centered(doc, text, size=12, bold=False, cn=CN_FONT, space_after=0, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_after = Pt(space_after); pf.space_before = Pt(space_before)
    pf.line_spacing = 1.5; pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    run = p.add_run(text)
    set_run_font(run, cn=cn, size=size, bold=bold)
    return p


def add_caption(doc, text, before=4, after=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = 1.2
    run = p.add_run(text)
    set_run_font(run, cn=CN_HEI, size=10.5, bold=True)
    return p


def add_image(doc, path, width_cm=14.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(6); pf.line_spacing = 1.2
    run = p.add_run(text)
    set_run_font(run, size=9, cn=CN_FONT)
    return p


def style_cell(cell, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_table(doc, headers, rows, col_align=None, header_size=10, body_size=10):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        style_cell(table.rows[0].cells[j], h, size=header_size, bold=True)
    for r in rows:
        cells = table.add_row().cells
        for j, val in enumerate(r):
            al = col_align[j] if col_align else WD_ALIGN_PARAGRAPH.CENTER
            style_cell(cells[j], val, size=body_size, align=al)
    return table


def add_toc(doc):
    add_centered(doc, "目  录", size=16, bold=True, cn=CN_HEI, space_after=12)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = 'TOC \\o "1-2" \\h \\z \\u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    run._element.append(f1); run._element.append(it); run._element.append(f2)
    ph = p.add_run("【请在 Word 中右键此处 →“更新域 / Update Field”自动生成目录】")
    set_run_font(ph, size=10.5, cn=CN_FONT, color=RGBColor(0x80, 0x80, 0x80))
    p2 = doc.add_paragraph(); r3 = p2.add_run()
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    r3._element.append(f3)


def enable_update_fields(doc):
    settings = doc.settings.element
    el = OxmlElement("w:updateFields"); el.set(qn("w:val"), "true")
    settings.append(el)


def add_footnotes_part(doc):
    """构建 footnotes.xml 部件并与文档关联。"""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    parts = [
        f'<w:footnote w:type="separator" w:id="-1"><w:p><w:pPr>'
        f'<w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
        f'<w:r><w:separator/></w:r></w:p></w:footnote>',
        f'<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:pPr>'
        f'<w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
        f'<w:r><w:continuationSeparator/></w:r></w:p></w:footnote>',
    ]
    for i, t in enumerate(_FOOTNOTES, start=1):
        t = t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        parts.append(
            f'<w:footnote w:id="{i}"><w:p>'
            f'<w:pPr><w:pStyle w:val="FootnoteText"/><w:spacing w:line="240" w:lineRule="auto"/></w:pPr>'
            f'<w:r><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteRef/></w:r>'
            f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" '
            f'w:eastAsia="{CN_FONT}"/><w:sz w:val="18"/></w:rPr>'
            f'<w:t xml:space="preserve"> {t}</w:t></w:r></w:p></w:footnote>'
        )
    xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        f'<w:footnotes xmlns:w="{W}">' + "".join(parts) + "</w:footnotes>"
    )
    partname = PackURI("/word/footnotes.xml")
    part = Part(partname, FOOTNOTES_CT, xml.encode("utf-8"), doc.part.package)
    doc.part.relate_to(part, FOOTNOTES_RT)


def add_page_number_footer(doc):
    p = doc.sections[-1].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    run._element.append(f1); run._element.append(it); run._element.append(f2)
    set_run_font(run, size=10.5)


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.0); sec.right_margin = Cm(3.0)

    style = doc.styles["Normal"]
    style.font.name = EN_FONT; style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    # ===================== 封面 =====================
    for _ in range(2):
        add_centered(doc, "")
    add_centered(doc, "对外经济贸易大学", size=22, bold=True, cn=CN_HEI, space_after=6)
    add_centered(doc, "《世界贸易组织概论》期末论文", size=18, bold=True, cn=CN_HEI, space_after=30)
    for _ in range(2):
        add_centered(doc, "")
    add_centered(doc, "入世是\u201c污染天堂\u201d还是\u201c技术外溢\u201d？", size=16, bold=True, cn=CN_HEI, space_after=4)
    add_centered(doc, "——基于跨国面板数据对 WTO 成员身份与出口 CO\u2082 强度关系的实证检验",
                 size=13, bold=True, cn=CN_HEI, space_after=40)
    for _ in range(3):
        add_centered(doc, "")
    info = [
        ("教学单位", "经贸学院"), ("课程名称", "世界贸易组织概论"),
        ("课程代码", "ITR305"), ("课序号", "2"),
        ("任课教师", "吕越、宁静馨"), ("考试方式", "论文"),
        ("学生姓名", "＿＿＿＿＿＿＿＿＿＿"), ("学    号", "＿＿＿＿＿＿＿＿＿＿"),
        ("学    期", "2026—2027 学年度第一学期"),
    ]
    for k, v in info:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.6; p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{k}："); set_run_font(r1, size=14, bold=True, cn=CN_HEI)
        r2 = p.add_run(v); set_run_font(r2, size=14)
    doc.add_page_break()

    # ===================== 摘要 =====================
    add_centered(doc, "入世是\u201c污染天堂\u201d还是\u201c技术外溢\u201d？", size=15, bold=True, cn=CN_HEI, space_after=4)
    add_centered(doc, "——基于跨国面板数据对 WTO 成员身份与出口 CO\u2082 强度关系的实证检验",
                 size=11, bold=True, cn=CN_HEI, space_after=10)
    add_centered(doc, "摘  要", size=12.5, bold=True, cn=CN_HEI, space_after=4)
    abstract = (
        "加入世界贸易组织（WTO）通过削减关税、约束壁垒与稳定贸易预期，深刻重塑了成员的出口结构，"
        "由此引发长期争论：贸易自由化究竟会使发展中国家沦为高碳产业的\u201c污染天堂\u201d，"
        "还是通过竞争与技术外溢促使出口部门\u201c变绿\u201d？本文以 2000—2020 年覆盖约 160 个经济体、"
        "共 3480 个国家—年观测的面板数据为样本，以出口碳强度为被解释变量，构建双向固定效应（TWFE）模型，"
        "并辅以事件研究、异质性分析与安慰剂检验。研究发现：第一，由于约 126 个创始成员在样本期内成员身份恒为 1，"
        "识别实际仅由 2001 年后加入的约 20 个经济体驱动，属交错型双重差分；第二，静态 TWFE 系数为 0.111（p=0.283）不显著，"
        "而事件研究显示入世后碳强度呈逐步累积的弱负向效应、部分后期系数边际显著，二者的背离恰反映了 TWFE 在交错处理下的局限；"
        "第三，碳强度的普遍下降主要由收入增长（EKC）、能效改善与结构调整驱动，异质性方向上提示\u201c污染天堂\u201d与\u201c技术外溢\u201d并存，"
        "但因处理组样本小而缺乏统计功效。结论是：WTO 成员身份本身并非碳强度的决定性因素，并据此提出\u201c贸易—环境—技术\u201d协同建议。"
    )
    add_body(doc, abstract, size=11)
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Pt(22)
    r1 = p.add_run("关键词："); set_run_font(r1, size=11, bold=True, cn=CN_HEI)
    r2 = p.add_run("世界贸易组织；出口碳强度；污染天堂假说；交错双重差分；碳边境调节机制")
    set_run_font(r2, size=11)
    doc.add_page_break()

    # ===================== 目录 =====================
    add_toc(doc)
    doc.add_page_break()

    # ===================== 引言 =====================
    add_h1(doc, "一、引言")
    add_body(doc,
        "自 1995 年世界贸易组织（WTO）成立以来，以最惠国待遇、国民待遇和关税约束为核心的多边贸易体制"
        "极大降低了跨国贸易成本，推动全球贸易额与生产网络迅速扩张。中国于 2001 年加入 WTO 后出口规模跃居世界首位，"
        "发展中国家在全球出口中的份额亦显著上升。然而，贸易扩张在带来增长红利的同时，也使\u201c贸易与环境\u201d的关系成为焦点。"
        "近年来欧盟碳边境调节机制（CBAM）落地，使\u201c贸易中隐含的碳\u201d从学术议题上升为现实的贸易政策工具，本文问题因而更具紧迫性。")
    add_body(doc,
        "围绕这一关系存在两种针锋相对的判断。一种是\u201c污染天堂假说\u201d：在环境规制存在国别差异时，"
        "贸易自由化会促使高碳生产环节向规制宽松的发展中国家转移，使其出口部门\u201c变脏\u201d、单位出口碳排放上升。"
        "另一种是\u201c技术外溢\u201d观点：融入多边体制会带来更激烈的竞争、更严格的进口国标准以及更便利的技术与资本流入，"
        "从而帮助出口企业提升能效、降低碳强度。两种机制方向相反，孰强孰弱本质上是有待数据检验的实证问题。")
    add_body(doc,
        "本文聚焦出口碳强度这一可度量、政策含义清晰的指标，检验 WTO 成员身份对它的因果影响。",
        footnote="出口碳强度 = CO\u2082 排放量 / 货物与服务出口额，单位为吨/百万美元，数值越低表示单位出口越清洁。")
    add_body(doc,
        "边际贡献在于：其一，直接以\u201c成员身份\u201d这一制度变量为处理变量，并明确将识别置于交错双重差分框架下讨论其来源与局限；"
        "其二，借助事件研究、安慰剂检验与异质性分析，缓解内生性并回应\u201c污染天堂\u201d与\u201c技术外溢\u201d之争。")

    # ===================== 文献综述 =====================
    add_h1(doc, "二、文献综述")
    add_body(doc,
        "贸易与环境关系的研究可追溯至 Grossman 和 Krueger（1991）的规模、结构与技术三效应框架，净方向取决于三者合力；"
        "环境库兹涅茨曲线（EKC）假说进一步认为污染随人均收入呈先升后降的倒 U 形。\u201c污染天堂\u201d文献中，"
        "Copeland 和 Taylor（1994、2004）从理论上论证了环境规制差异如何塑造污染密集型产业的国际分工，"
        "并指出其经验证据常因要素禀赋效应对冲而较弱；Levinson、Antweiler 等的实证结论分歧明显。"
        "与之相对，Frankel 和 Rose（2005）提供了贸易可改善环境质量的证据；就 WTO/GATT 的作用而言，Rose（2004）曾质疑其贸易促进效应。")
    add_body(doc,
        "方法论上，近年交错双重差分文献（Goodman-Bacon，2021；de Chaisemartin 和 D'Haultfœuille，2020；Sun 和 Abraham，2021）指出，"
        "当处理时点交错且效应异质时，静态双向固定效应估计会对不同时点的处理效应赋予可能为负的权重，并使用\u201c已处理单元\u201d作为对照，"
        "从而产生偏误。这一进展对本文识别\u201c入世\u201d这一交错制度冲击尤为关键，但既有相关实证研究多以贸易开放度为自变量、对此着墨不多，本文力图补充。")

    # ===================== 理论与假设 =====================
    add_h1(doc, "三、理论分析与研究假设")
    add_body(doc,
        "加入 WTO 对出口碳强度的影响可经三条路径传导：结构路径上，关税减让使一国按比较优势重配资源，"
        "若其优势位于碳密集产业，出口结构\u201c变脏\u201d、碳强度上升，对应\u201c污染天堂\u201d机制；"
        "技术路径上，进口中间品与外资的技术外溢、叠加出口目的国的环境标准，推动企业采用更清洁工艺、碳强度下降，对应\u201c技术外溢\u201d机制；"
        "规模路径则放大前两者作用。由于结构与技术效应方向相反，二者强弱取决于成员国的禀赋与发展阶段，据此提出竞争性假设：")
    add_body(doc, "H1a（污染天堂）：加入 WTO 后，成员出口碳强度上升；", first_indent=True)
    add_body(doc, "H1b（技术外溢）：加入 WTO 后，成员出口碳强度下降；", first_indent=True)
    add_body(doc,
        "H2（异质性）：能源强度高、处于工业化中期的发展中国家更可能呈\u201c污染天堂\u201d特征；"
        "能源强度低、技术水平高的经济体更可能呈\u201c技术外溢\u201d特征。", first_indent=True)

    # ===================== 数据与方法 =====================
    add_h1(doc, "四、数据来源与研究设计")
    add_h2(doc, "（一）数据与变量")
    add_body(doc,
        "本文构建 2000—2020 年跨国面板，覆盖约 160 个经济体，经缺失值处理后共 3480 个国家—年观测。"
        "被解释变量为 ln(CO\u2082 强度)；核心解释变量为 WTO 成员身份虚拟变量；控制变量包括人均 GDP 及其平方项（检验 EKC）、能源强度、工业增加值占比。",
        footnote="出口额取自世界银行 WDI 数据库，指标代码 NE.EXP.GNFS.CD（现价美元）；人均 GDP、能源强度与工业占比同样取自世界银行 WDI；CO\u2082 排放数据来自公开排放数据库（如 EDGAR / Our World in Data）。")
    add_caption(doc, "表 1  主要变量描述性统计")
    add_table(doc,
        ["变量", "观测值", "均值", "标准差", "最小值", "中位数", "最大值"],
        [
            ["CO\u2082强度（吨/百万美元）", "3480", "2246.74", "3123.27", "1.14", "1349.75", "45695.13"],
            ["ln(CO\u2082强度)", "3480", "7.137", "1.182", "0.127", "7.208", "10.730"],
            ["WTO成员（0/1）", "3480", "0.792", "0.406", "0", "1", "1"],
            ["CO\u2082排放量（百万吨）", "3480", "188.36", "833.32", "0.001", "16.03", "11964.13"],
            ["人均GDP（2015美元）", "3452", "14395.0", "19659.5", "287.4", "4993.6", "118382.9"],
            ["能源强度", "3418", "5.231", "3.213", "0.110", "4.330", "30.440"],
            ["工业占GDP比重（%）", "3392", "27.30", "12.35", "3.72", "25.03", "86.67"],
        ],
        col_align=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER]*6,
        body_size=9.5, header_size=9.5)
    add_note(doc, "注：CO\u2082 强度右偏明显，故取对数；WTO 成员观测占比 79.2%。")
    add_body(doc,
        "图 1a 比较处理组与控制组的 ln(CO\u2082 强度) 趋势，两组均呈下降；图 1b 显示其分布大致对称、均值约 7.14。"
        "图 2 进一步对比主要出口国 2000 年与 2020 年的碳强度，中国、印度、俄罗斯、南非等普遍下降六成以上，呈全局性\u201c变绿\u201d。")
    add_image(doc, "fig1.png")
    add_caption(doc, "图 1  描述性统计：处理组/控制组趋势与 ln(CO\u2082 强度) 分布")
    add_image(doc, "fig2.png")
    add_caption(doc, "图 2  主要国家出口 CO\u2082 强度：2000 年与 2020 年对比")

    add_h2(doc, "（二）计量模型与识别策略")
    add_body(doc,
        "为识别成员身份的净效应，设定双向固定效应模型：")
    add_body(doc,
        "ln(CO\u2082 强度)_{it} = β\u2081·WTO_{it} + β\u2082·ln(pgdp)_{it} + β\u2083·[ln(pgdp)]\u00b2_{it} "
        "+ β\u2084·能源强度_{it} + β\u2085·工业占比_{it} + μ_i + λ_t + ε_{it}",
        first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_body(doc,
        "其中 μ_i、λ_t 分别为国家与年份固定效应，β\u2081 为所关注效应，标准误在国家层面聚类。")
    add_body(doc,
        "需要特别说明识别来源：样本中约 126 个 1995 年 WTO/GATT 创始成员在整个区间内成员身份恒为 1，不提供组内时间变异；"
        "真正驱动 β\u2081 的是 2001 年后陆续加入的约 20 个经济体，因此本质上是一个交错型双重差分（staggered DID）。",
        footnote="主要时间变异来自中国（2001）、沙特阿拉伯（2005）、越南（2007）、乌克兰（2008）、俄罗斯（2012）等。")
    add_body(doc,
        "这一结构带来两点必须正视的局限：其一，这些处理国高度异质——中国为制造业出口大国，俄罗斯、沙特为能源出口国，越南为新兴加工贸易国，"
        "将其合并为单一\u201c处理组\u201d与创始成员\u201c控制组\u201d比较，可比性较弱；其二，处理单元仅约 20 个，统计功效有限。"
        "故本文以事件研究为主要识别工具，以静态 TWFE 系数为辅，并对二者差异保持审慎。")

    # ===================== 实证结果 =====================
    add_h1(doc, "五、实证结果与讨论")
    add_h2(doc, "（一）基准回归")
    add_body(doc,
        "表 2 报告基准回归。混合 OLS 中 WTO 系数为 -0.318（p<0.05），但该设定未控制国家异质性，存在严重选择性偏误"
        "——本身碳强度较低的发达经济体往往更早成为成员。加入国家固定效应后系数降至 -0.025 且不显著；"
        "进一步加入年份固定效应（双向 FE）后，系数为 0.111、p=0.283，不显著。控制变量与理论高度吻合："
        "ln(人均 GDP) 显著为正、其平方项显著为负，刻画出倒 U 形 EKC；能源强度显著为正，工业占比显著为负，"
        "说明碳强度的下降主要由收入增长、能效提升与结构调整解释，而非入世这一制度事件本身。")
    add_caption(doc, "表 2  基准回归结果（被解释变量：ln(CO\u2082 强度)）")
    add_table(doc,
        ["变量", "(1) 混合OLS", "(2) 国家FE", "(3) 双向FE"],
        [
            ["WTO成员（β\u2081）", "-0.3181**\n(0.1388)", "-0.0251\n(0.1333)", "0.1108\n(0.1031)"],
            ["ln(人均GDP)", "2.9522***\n(0.5117)", "2.1157**\n(0.8535)", "2.9030***\n(0.7438)"],
            ["[ln(人均GDP)]\u00b2", "-0.1861***\n(0.0289)", "-0.2085***\n(0.0501)", "-0.2072***\n(0.0427)"],
            ["能源强度", "0.0968***\n(0.0231)", "0.0470***\n(0.0135)", "0.0433***\n(0.0148)"],
            ["工业增加值占比", "0.0018\n(0.0049)", "-0.0111***\n(0.0041)", "-0.0173***\n(0.0030)"],
            ["国家固定效应", "否", "是", "是"],
            ["年份固定效应", "否", "否", "是"],
        ],
        col_align=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER]*3,
        body_size=9.5, header_size=10)
    add_note(doc, "注：括号内为国家层面聚类稳健标准误；*、**、*** 分别表示在 10%、5%、1% 水平显著。")

    add_h2(doc, "（二）平行趋势与事件研究")
    add_body(doc,
        "事件研究以入世前一年（t-1）为基准。为更严格地检验平行趋势，本文对入世前各期（t-6 至 t-2）系数做联合显著性检验，"
        "不能拒绝其联合为零的原假设（Wald 检验：F=1.07，p=0.38），且图 3 中入世前系数及其 95% 置信区间均覆盖零，"
        "为平行趋势假定提供了较正式的支持。入世后效应随时间逐步累积：t0、t+1 期不显著，t+2 期起部分系数在 5% 水平上边际显著为负"
        "（t+2 期 β≈-0.13、p≈0.04，t+6 期 β≈-0.27、p≈0.03），呈现\u201c渐进累积但统计不稳定\u201d的形态。")
    add_image(doc, "fig3.png")
    add_caption(doc, "图 3  事件研究：WTO 入世对出口 CO\u2082 强度的动态影响")
    add_body(doc,
        "事件研究的负向动态与静态 TWFE 的不显著正系数看似矛盾，实则可调和，原因有三。其一，静态 β\u2081 是对所有入世后年份的加权平均，"
        "而早期（t0、t+1）效应接近零，会稀释后期较强的负向效应，使平均值偏向零。其二，更关键地，在交错处理与效应异质的情形下，"
        "静态 TWFE 会使用\u201c较早入世国\u201d作为\u201c较晚入世国\u201d的对照（即\u201c坏比较\u201d），并对部分 ATT 赋予负权重，"
        "从而使点估计的符号与幅度失真甚至变号（Goodman-Bacon，2021）；以干净的 t-1 为基准的事件研究估计更为可信。"
        "其三，处理组仅约 20 国，动态系数的显著性对样本与设定较敏感，故应将其解读为\u201c存在微弱、渐进的降碳迹象\u201d，"
        "而非稳健结论。综合而言，更稳妥的判断是：入世未带来强而稳健的碳强度变化，证据方向偏弱负但高度不确定。")

    add_h2(doc, "（三）异质性分析")
    add_body(doc,
        "图 4 显示，以发达国家为基准，高能源强度国家效应约 +15.7%、发展中国家约 +8.3%（方向符合\u201c污染天堂\u201d），"
        "低能源强度国家约 -4.7%（方向符合\u201c技术外溢\u201d）。但须强调：各组系数 p 值均大于 0.3、置信区间普遍跨越零，"
        "而\u201c不显著\u201d并不等于\u201c存在异质性\u201d——它更可能反映处理组仅约 20 国所致的统计功效不足。"
        "严格而言，应通过\u201cWTO×分组\u201d交互项并对组间系数相等做正式检验；本文数据无法拒绝组间相等，故异质性只能作为方向性、提示性证据。")
    add_image(doc, "fig4.png")
    add_caption(doc, "图 4  WTO 效应的异质性分析")

    add_h2(doc, "（四）稳健性、安慰剂与测度问题")
    add_body(doc,
        "表 3 显示，剔除中国后系数为 0.113（p=0.279），1% 缩尾后为 0.122（p=0.216），均与基准一致。"
        "图 5 的安慰剂检验将\u201c入世年份\u201d随机赋值并重复 500 次，所得系数分布大致以零为中心，真实估计值 0.111 恰落入分布主体，"
        "说明静态正系数与随机噪声难以区分。")
    add_caption(doc, "表 3  稳健性检验")
    add_table(doc,
        ["检验方案", "β 系数", "效应 (%)", "p 值"],
        [
            ["基准模型（双向FE）", "0.1108", "11.72", "0.283"],
            ["排除中国", "0.1134", "12.01", "0.279"],
            ["去除极端值（1% 缩尾）", "0.1220", "12.98", "0.216"],
        ],
        col_align=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER]*3,
        body_size=10, header_size=10)
    add_image(doc, "fig5.png")
    add_caption(doc, "图 5  安慰剂检验：随机入世日期的系数分布")
    add_body(doc,
        "此外，被解释变量本身存在一个值得警惕的测度问题：出口碳强度以出口额为分母，而入世往往会显著扩大出口额（分母）。"
        "若出口激增而排放未同步变化，碳强度会\u201c机械地\u201d下降，从而高估\u201c技术外溢\u201d。这意味着任何负向估计都可能部分源于分母效应，"
        "而本文在控制收入与结构后仍未发现稳健的负向平均效应，恰使\u201c入世并未显著降碳\u201d的结论更趋保守。"
        "更彻底的做法是分别考察分子（排放）与分母（出口）的反应，或改用投入产出法测度出口\u201c隐含碳\u201d，本文受数据所限留待后续。")

    # ===================== 结论 =====================
    add_h1(doc, "六、结论与政策启示")
    add_body(doc,
        "本文利用 2000—2020 年跨国面板数据检验了 WTO 成员身份对出口碳强度的影响，得到三点结论。第一，识别仅由约 20 个交错入世国驱动，"
        "静态 TWFE 显示无显著平均效应，事件研究则显示微弱、渐进的负向动态，二者背离主要源于交错处理下 TWFE 的偏误与处理组样本过小。"
        "第二，碳强度的普遍下降主要由收入增长（EKC）、能效改善与结构调整驱动，而非入世本身。"
        "第三，异质性方向上提示\u201c污染天堂\u201d与\u201c技术外溢\u201d并存，但均因功效不足而不显著。")
    add_body(doc,
        "政策含义清晰：其一，不应简单地将 WTO 成员身份或贸易自由化视为碳强度上升的根源，贸易的环境效应取决于成员国内部的禀赋与配套政策；"
        "其二，能源密集型发展中国家应警惕高碳产业转入风险，通过提高能效标准、完善环境规制与碳定价，避免比较优势固化于高碳环节；"
        "其三，应发挥多边体制在技术扩散方面的正向作用，推动绿色技术与环境产品自由化，并审慎设计 CBAM 等碳边境措施，"
        "使其既防止碳泄漏，又不至于演变为变相的绿色保护主义。")
    add_body(doc,
        "本文局限亦明显：处理组样本小、统计功效有限；以国别加总数据与二值成员身份衡量，难以剥离行业构成、价值链分工与关税约束深度的差异；"
        "出口碳强度的分母效应可能干扰解读。未来可采用更稳健的交错 DID 估计量（如 Sun-Abraham），并结合行业、企业层面与投入产出\u201c隐含碳\u201d数据，"
        "进一步识别贸易自由化影响碳强度的具体渠道。")

    # ===================== 参考文献 =====================
    add_h1(doc, "参考文献")
    refs = [
        "[1] Grossman, G. M., & Krueger, A. B. Environmental Impacts of a North American Free Trade Agreement[R]. NBER Working Paper No. 3914, 1991.",
        "[2] Copeland, B. R., & Taylor, M. S. North-South Trade and the Environment[J]. The Quarterly Journal of Economics, 1994, 109(3): 755-787.",
        "[3] Copeland, B. R., & Taylor, M. S. Trade, Growth, and the Environment[J]. Journal of Economic Literature, 2004, 42(1): 7-71.",
        "[4] Antweiler, W., Copeland, B. R., & Taylor, M. S. Is Free Trade Good for the Environment?[J]. American Economic Review, 2001, 91(4): 877-908.",
        "[5] Frankel, J. A., & Rose, A. K. Is Trade Good or Bad for the Environment? Sorting Out the Causality[J]. The Review of Economics and Statistics, 2005, 87(1): 85-91.",
        "[6] Rose, A. K. Do We Really Know That the WTO Increases Trade?[J]. American Economic Review, 2004, 94(1): 98-114.",
        "[7] Levinson, A., & Taylor, M. S. Unmasking the Pollution Haven Effect[J]. International Economic Review, 2008, 49(1): 223-254.",
        "[8] Goodman-Bacon, A. Difference-in-Differences with Variation in Treatment Timing[J]. Journal of Econometrics, 2021, 225(2): 254-277.",
        "[9] de Chaisemartin, C., & D'Haultf\u0153uille, X. Two-Way Fixed Effects Estimators with Heterogeneous Treatment Effects[J]. American Economic Review, 2020, 110(9): 2964-2996.",
        "[10] Sun, L., & Abraham, S. Estimating Dynamic Treatment Effects in Event Studies with Heterogeneous Treatment Effects[J]. Journal of Econometrics, 2021, 225(2): 175-199.",
        "[11] World Bank. World Development Indicators[DB/OL]. Washington, D.C.: The World Bank, 2024.",
        "[12] 李小平, 卢现祥. 国际贸易、污染产业转移和中国工业 CO\u2082 排放[J]. 经济研究, 2010(1): 15-26.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.5; pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.space_after = Pt(2); pf.left_indent = Pt(24); pf.first_line_indent = Pt(-24)
        run = p.add_run(r); set_run_font(run, size=10.5)

    add_page_number_footer(doc)
    add_footnotes_part(doc)
    enable_update_fields(doc)
    out = "WTO期末论文_入世与出口碳强度.docx"
    doc.save(out)
    print("已生成：", out, "｜脚注数：", len(_FOOTNOTES))


if __name__ == "__main__":
    missing = [f for f in ["fig1.png","fig2.png","fig3.png","fig4.png","fig5.png"] if not os.path.exists(f)]
    if missing:
        raise SystemExit("缺少图片，请先运行 generate_figures.py：" + ", ".join(missing))
    main()
